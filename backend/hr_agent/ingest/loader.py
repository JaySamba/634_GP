"""
Document loader.

Walks the HR documents folder and returns a Document object per file.
Handles PDFs (text-extracted) and DOCX. Skips files that aren't readable.

The corpus inspector confirmed all 142 files are well-formed, so this
keeps logic minimal — no scanned PDF / OCR handling here.
"""

from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterator

from pypdf import PdfReader
from docx import Document as DocxDocument


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}

# Files the inspector writes — we never want to ingest them
EXCLUDE_FILENAMES = {"corpus_report.json", "corpus_report.md"}


@dataclass
class Document:
    """One source document loaded from disk."""

    filename: str           # e.g. "HR-PP-50B AODA Policy.docx"
    path: str               # absolute path
    text: str               # full extracted text
    file_type: str          # "pdf" | "docx" | "text"

    # Optional structural info — populated where we can extract it
    page_count: int = 0
    paragraphs: list[str] = field(default_factory=list)


def _load_pdf(path: Path) -> Document:
    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    text = "\n\n".join(pages)
    return Document(
        filename=path.name,
        path=str(path),
        text=text,
        file_type="pdf",
        page_count=len(reader.pages),
    )


def _load_docx(path: Path) -> Document:
    doc = DocxDocument(str(path))

    # Collect paragraphs — keep them as a list, the chunker uses headings
    paragraphs = []
    for p in doc.paragraphs:
        if p.text.strip():
            # Tag heading paragraphs so the chunker can split on them
            style = (p.style.name if p.style else "") or ""
            if style.startswith("Heading"):
                paragraphs.append(f"[HEADING] {p.text}")
            else:
                paragraphs.append(p.text)

    # Also pull table content — HR docs use tables a lot
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return Document(
        filename=path.name,
        path=str(path),
        text="\n".join(paragraphs),
        file_type="docx",
        paragraphs=paragraphs,
    )


def _load_text(path: Path) -> Document:
    return Document(
        filename=path.name,
        path=str(path),
        text=path.read_text(encoding="utf-8", errors="replace"),
        file_type="text",
    )


def load_document(path: Path) -> Document | None:
    """Load a single file. Returns None if unsupported or unreadable."""
    ext = path.suffix.lower()
    try:
        if ext == ".pdf":
            return _load_pdf(path)
        if ext == ".docx":
            return _load_docx(path)
        if ext in (".txt", ".md"):
            return _load_text(path)
    except Exception as e:
        print(f"  ⚠  skipped {path.name}: {type(e).__name__}: {e}")
        return None
    return None


def iter_documents(folder: Path) -> Iterator[Document]:
    """Yield Document objects for every supported file in the folder."""
    folder = folder.expanduser().resolve()
    if not folder.is_dir():
        raise NotADirectoryError(f"{folder} is not a directory")

    files = sorted(
        f for f in folder.rglob("*")
        if f.is_file()
        and f.suffix.lower() in SUPPORTED_EXTENSIONS
        and f.name not in EXCLUDE_FILENAMES
    )

    for path in files:
        doc = load_document(path)
        if doc and doc.text.strip():
            yield doc
